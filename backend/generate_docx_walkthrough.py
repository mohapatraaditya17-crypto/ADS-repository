import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_element(name):
    return OxmlElement(name)

def add_p_border_bottom(p, color_hex="CCCCCC", size=12):
    pPr = p._p.get_or_add_pPr()
    pBdr = create_element('w:pBdr')
    bottom = create_element('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_docx():
    md_path = r"c:\Users\us183046\OneDrive - Grant Thornton Advisors LLC\Desktop\Falcon LLM\walkthrough.md"
    docx_path = r"c:\Users\us183046\OneDrive - Grant Thornton Advisors LLC\Desktop\Falcon LLM\walkthrough.docx"
    
    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found")
        return
        
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    doc = Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Document Title Style
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    lines = content.split('\n')
    in_code_block = False
    code_lines = []
    
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End of code block, write the block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.right_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                
                # Add shading to code block (using a 1x1 table)
                tbl = doc.add_table(rows=1, cols=1)
                tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
                tbl.autofit = False
                tbl.columns[0].width = Inches(5.5)
                cell = tbl.cell(0, 0)
                set_cell_background(cell, "F4F4F6")
                
                # Add code text
                cell_p = cell.paragraphs[0]
                cell_p.paragraph_format.space_before = Pt(4)
                cell_p.paragraph_format.space_after = Pt(4)
                code_text = "\n".join(code_lines)
                run = cell_p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x30, 0x30, 0x30)
                
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
            i += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
            
        # Parse titles
        if line.startswith('# '):
            title_text = line[2:].strip()
            p = doc.add_heading(level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(title_text)
            run.font.name = 'Arial'
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0C, 0x23, 0x40) # Deep Navy
            
            # Bottom border
            add_p_border_bottom(p, color_hex="0C2340", size=18)
            
        elif line.startswith('## '):
            heading_text = line[3:].strip()
            p = doc.add_heading(level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(heading_text)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1D, 0x70, 0xB8) # Electric Blue
            
        elif line.startswith('### '):
            heading_text = line[4:].strip()
            p = doc.add_heading(level=2)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(heading_text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x5A, 0x9C)
            
        elif line.startswith('---'):
            # Horizontal rule
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_p_border_bottom(p, color_hex="DDDDDD", size=6)
            
        elif line.strip().startswith(('* ', '- ')):
            # List item
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.25)
            
            parse_formatted_text(p, text)
            
        elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            # Numbered list item
            text = line.strip().split('.', 1)[1].strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.25)
            
            parse_formatted_text(p, text)
            
        elif line.startswith('    * ') or line.startswith('    - ') or line.startswith('  * ') or line.startswith('  - '):
            # Nested list item
            text = line.strip()[2:].strip()
            p = doc.add_paragraph(style='List Bullet 2')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.5)
            
            parse_formatted_text(p, text)
            
        elif line.strip():
            # Paragraph
            text = line.strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
            parse_formatted_text(p, text)
            
        i += 1
        
    doc.save(docx_path)
    print(f"Success! Saved docx to {docx_path}")

def parse_formatted_text(paragraph, text):
    # Regex splits by bold, code, or links
    parts = re.split(r'(\*\*.*?\*\*|`.*?`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
            
        if part.startswith('**') and part.endswith('**'):
            # Bold
            val = part[2:-2]
            run = paragraph.add_run(val)
            run.font.bold = True
        elif part.startswith('`') and part.endswith('`'):
            # Code
            val = part[1:-1]
            run = paragraph.add_run(val)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00) # Dark red for inline code
        elif part.startswith('[') and ']' in part and '(' in part and part.endswith(')'):
            # Link
            match = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if match:
                link_text, link_url = match.groups()
                run = paragraph.add_run(link_text)
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xEE)
                run.font.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    build_docx()
