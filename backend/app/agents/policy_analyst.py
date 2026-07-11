import logging
import time
import os
import csv
from typing import Generator, List, Dict, Any
from app.tools.crowdstrike.policies import (
    list_prevention_policies,
    list_sensor_update_policies,
    list_device_control_policies,
    list_firewall_policies,
    list_host_groups,
    get_ml_exclusions,
    get_ioa_exclusions,
    get_sve_exclusions,
)
from app.tools.crowdstrike.hosts import list_all_hosts
from app.agents.llm import call_llm_stream
from docx import Document

logger = logging.getLogger("policy_analyst")

SYSTEM_PROMPT = """
You are the "Policy Analyst" agent for the Falcon AI Copilot. Your role is to examine security policies (Prevention, Sensor Update, Device Control, Firewall, Exclusions) across the environment, identify coverage gaps, compare rules, and recommend best practices.

CRITICAL INSTRUCTIONS:
1. Ground your recommendations directly in standard CrowdStrike security best practices (e.g. recommending "Moderate" or "Aggressive" machine learning threshold configurations for prevention).
2. For comparisons, present settings side-by-side using Markdown tables showing Settings | Policy A | Policy B.
3. If the user asks to list ANY type of policies or exclusions (Prevention, Device Control, Firewall, Sensor Update, ML/IOA/SVE Exclusions), you MUST present them in a Markdown table. DO NOT output a bulleted list.
4. Call out disabled protections (e.g., Script Execution monitoring disabled) as Gaps or Warnings.
5. Falcon AI Copilot is strictly READ-ONLY. Emphasize that you generate hardening recommendations, but DO NOT apply modifications directly.
6. STRICT ANTI-HALLUCINATION RULE: If you are unable to answer a question or if the API returns an error or no data, DO NOT make up or hallucinate any data. You MUST return a clear statement explaining the reason (e.g., "couldn't reach server", "unable to verify", "API error", "no records found").
7. If the user asks to list host groups (or groups), you MUST present the host groups in a Markdown table with the exact columns: Group Name | Group ID | Hosts Applied | Group Type | Group Description | Created By | Created Date | Last Modified By | Last Modified | Assignment Rule. Use 'NA' for any missing or unretrieved information. Do not display a policies table for this request. Do NOT output a bulleted list.
8. LASER-FOCUS RULE: Provide answers strictly aligned with the specific keywords and intent of the user's query. Do not provide generalized overviews, unsolicited tangential information, or verbose preambles. Keep responses concise and directly address the exact ask.
9. If the context contains a 'Generated Report' or 'Generated CSV' markdown link, you MUST output the link verbatim in your response. The system has already generated the file! DO NOT claim you are an AI and cannot generate files or download files. You MUST simply present the link to the user.
10. In any Markdown tables you generate, represent Boolean 'enabled' fields as 'Enabled' or 'Disabled' instead of 'True' or 'False'.
11. Use the exact description provided in the context for policies. Do not alter or summarize it. If no description is available in the context (empty or missing), you MUST output 'NA' in the table description column.
"""

def _compress_policy(policy: Dict[str, Any], include_settings: bool = False) -> Dict[str, Any]:
    """Compresses a prevention policy dictionary to save context space."""
    compressed = {
        "id": policy.get("id"),
        "name": policy.get("name"),
        "platform_name": policy.get("platform_name"),
        "enabled": "Enabled" if policy.get("enabled") else "Disabled",
        "description": policy.get("description") or "NA",
        "groups": [g.get("name") for g in policy.get("groups", []) if isinstance(g, dict) and "name" in g]
    }
    
    if include_settings:
        settings_summary = {}
        for category in policy.get("prevention_settings", []):
            cat_name = category.get("name")
            cat_settings = {}
            for setting in category.get("settings", []):
                s_id = setting.get("id")
                s_val = setting.get("value", {})
                enabled = s_val.get("enabled", False)
                configured = s_val.get("configured", False)
                cat_settings[s_id] = f"enabled={enabled}" if not configured else f"enabled={enabled}(configured)"
            if cat_settings:
                settings_summary[cat_name] = cat_settings
        compressed["prevention_settings_summary"] = settings_summary
        
    return compressed

def _compress_generic_policy(policy: Dict[str, Any]) -> Dict[str, Any]:
    """Compresses a generic policy dictionary to basic fields."""
    return {
        "id": policy.get("id"),
        "name": policy.get("name"),
        "platform_name": policy.get("platform_name"),
        "enabled": "Enabled" if policy.get("enabled") else "Disabled",
        "description": policy.get("description") or "NA",
        "groups": [g.get("name") for g in policy.get("groups", []) if isinstance(g, dict) and "name" in g]
    }

def _compress_host_group(group: Dict[str, Any]) -> Dict[str, Any]:
    """Compresses a host group dictionary to include detailed auditing fields."""
    return {
        "id": group.get("id", "NA") or "NA",
        "name": group.get("name", "NA") or "NA",
        "hosts_applied": group.get("hosts_applied", "NA"),
        "group_type": group.get("group_type", "NA") or "NA",
        "description": group.get("description", "NA") or "NA",
        "created_by": group.get("created_by", "NA") or "NA",
        "created_timestamp": group.get("created_timestamp", "NA") or "NA",
        "modified_by": group.get("modified_by", "NA") or "NA",
        "modified_timestamp": group.get("modified_timestamp", "NA") or "NA",
        "assignment_rule": group.get("assignment_rule", "NA") or "NA"
    }

def run_policy_analyst(
    user_query: str, 
    chat_history: List[Dict[str, str]] = []
) -> Generator[Dict[str, Any], None, None]:
    """
    Policy Analyst Agent runner. Gathers policies, runs reviews, and streams text tokens.
    """
    logger.info(f"Running Policy Analyst for query: '{user_query}'")
    q = user_query.lower()
    
    platform = None
    if "windows" in q:
        platform = "Windows"
    elif "linux" in q:
        platform = "Linux"
    elif "mac" in q:
        platform = "Mac"
    
    # 1. Loading policies notification
    yield {
        "event": "agent_state",
        "data": {
            "agent": "Policy Analyst",
            "message": "Retrieving CrowdStrike security configurations..."
        }
    }
    
    fetch_device = "device" in q or "usb" in q
    fetch_firewall = "firewall" in q
    fetch_update = "update" in q or "version" in q or "sensor" in q
    fetch_prev = "prevention" in q or "protection" in q or "malware" in q or "quarantine" in q
    fetch_exclusions = "exclusion" in q or "exclude" in q
    fetch_host_groups = "group" in q
    
    generate_docx = "docx" in q or "word" in q or "export" in q or "download" in q
    generate_csv = "csv" in q
    
    # If they just asked for "policies" but didn't specify which one, default to prevention or all?
    if not any([fetch_device, fetch_firewall, fetch_update, fetch_prev, fetch_exclusions, fetch_host_groups]) and "policy" in q:
        fetch_prev = True # Default to prevention policies if unspecified
    
    prev_policies = []
    if fetch_prev:
        yield {
            "event": "tool_start",
            "data": {
                "name": "list_prevention_policies",
                "params": {"platform": platform} if platform else {}
            }
        }
        start_time = time.time()
        prev_policies = list_prevention_policies(platform=platform)
        yield {
            "event": "tool_complete",
            "data": {
                "name": "list_prevention_policies",
                "status": "success",
                "duration_ms": int((time.time() - start_time) * 1000),
                "result": {"policies_retrieved": len(prev_policies)}
            }
        }
    
    other_policies = {}
    if fetch_firewall:
        other_policies["firewall"] = list_firewall_policies()
    if fetch_device:
        other_policies["device_control"] = list_device_control_policies()
    if fetch_update:
        other_policies["sensor_update"] = list_sensor_update_policies()
        
    exclusions = {}
    if fetch_exclusions:
        if "ml" in q or "machine learning" in q:
            exclusions["ml"] = get_ml_exclusions()
        if "ioa" in q or "indicator" in q:
            exclusions["ioa"] = get_ioa_exclusions()
        if "sve" in q or "sensor visibility" in q:
            exclusions["sve"] = get_sve_exclusions()
        # If no specific exclusion was mentioned, get all
        if not exclusions:
            exclusions["ml"] = get_ml_exclusions()
            exclusions["ioa"] = get_ioa_exclusions()
            exclusions["sve"] = get_sve_exclusions()
        
    host_groups = []
    hosts = []
    if fetch_host_groups:
        yield {
            "event": "tool_start",
            "data": {
                "name": "list_host_groups",
                "params": {}
            }
        }
        t0 = time.time()
        host_groups = list_host_groups()
        yield {
            "event": "tool_complete",
            "data": {
                "name": "list_host_groups",
                "status": "success",
                "duration_ms": int((time.time() - t0) * 1000),
                "result": {"groups_retrieved": len(host_groups)}
            }
        }
        
        # If exporting to DOCX or CSV for groups, they usually want hosts as well
        if generate_docx or generate_csv:
            yield {
                "event": "tool_start",
                "data": {
                    "name": "list_all_hosts",
                    "params": {"max_results": 1000}
                }
            }
            t0 = time.time()
            try:
                hosts = list_all_hosts(max_results=1000)
            except Exception as e:
                logger.error(f"Failed to list all hosts: {e}")
                hosts = []
                
            yield {
                "event": "tool_complete",
                "data": {
                    "name": "list_all_hosts",
                    "status": "success",
                    "duration_ms": int((time.time() - t0) * 1000),
                    "result": {"hosts_retrieved": len(hosts or [])}
                }
            }
            
        
    yield {
        "event": "agent_state",
        "data": {
            "agent": "Policy Analyst",
            "message": "Analyzing policies and compiling gap reviews..."
        }
    }
    
    # Build prompt context
    need_settings = any(k in q for k in ["compare", "gap", "review", "recommend", "setting", "configure", "tampering", "visibility", "ml", "malware", "quarantine"])
    compressed_prev = [_compress_policy(p, include_settings=need_settings) for p in prev_policies]
    
    compressed_others = {}
    for k, p_list in other_policies.items():
        compressed_others[k] = [_compress_generic_policy(p) for p in p_list]

    prompt_context = ""
    if prev_policies:
        prompt_context += f"Prevention Policies Configuration:\n{compressed_prev}\n"
    if host_groups:
        compressed_groups = [_compress_host_group(g) for g in host_groups]
        prompt_context += f"\nHost Groups Configuration:\n{compressed_groups}\n"
    if other_policies:
        prompt_context += f"\nOther Policies Configurations:\n{compressed_others}\n"
    if exclusions:
        prompt_context += f"\nExclusions Configurations:\n{exclusions}\n"
        
    download_links = []
    if generate_docx and fetch_host_groups:
        try:
            yield {
                "event": "agent_state",
                "data": {
                    "agent": "Policy Analyst",
                    "message": "Generating DOCX report for Hosts and Host Groups..."
                }
            }
            
            doc = Document()
            doc.add_heading('Falcon AI Copilot - Hosts and Host Groups', 0)
            
            # Host Groups Table
            doc.add_heading('Host Groups', level=1)
            if host_groups:
                table = doc.add_table(rows=1, cols=10)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Group Name'
                hdr_cells[1].text = 'Group ID'
                hdr_cells[2].text = 'Hosts Applied'
                hdr_cells[3].text = 'Group Type'
                hdr_cells[4].text = 'Group Description'
                hdr_cells[5].text = 'Created By'
                hdr_cells[6].text = 'Created Date'
                hdr_cells[7].text = 'Last Modified By'
                hdr_cells[8].text = 'Last Modified'
                hdr_cells[9].text = 'Assignment Rule'
                for g in host_groups:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(g.get("name", "NA") or "NA")
                    row_cells[1].text = str(g.get("id", "NA") or "NA")
                    row_cells[2].text = str(g.get("hosts_applied", "NA"))
                    row_cells[3].text = str(g.get("group_type", "NA") or "NA")
                    row_cells[4].text = str(g.get("description", "NA") or "NA")
                    row_cells[5].text = str(g.get("created_by", "NA") or "NA")
                    row_cells[6].text = str(g.get("created_timestamp", "NA") or "NA")
                    row_cells[7].text = str(g.get("modified_by", "NA") or "NA")
                    row_cells[8].text = str(g.get("modified_timestamp", "NA") or "NA")
                    row_cells[9].text = str(g.get("assignment_rule", "NA") or "NA")
            else:
                doc.add_paragraph("No host groups found.")
                
            # Hosts Table
            doc.add_heading('Managed Hosts', level=1)
            if hosts:
                htable = doc.add_table(rows=1, cols=4)
                htable.style = 'Table Grid'
                h_hdr = htable.rows[0].cells
                h_hdr[0].text = 'Hostname'
                h_hdr[1].text = 'OS Version'
                h_hdr[2].text = 'Local IP'
                h_hdr[3].text = 'Status'
                for h in (hosts or []):
                    h_row = htable.add_row().cells
                    h_row[0].text = str(h.get("hostname", "NA") or "NA")
                    h_row[1].text = str(h.get("os_version", "NA") or "NA")
                    h_row[2].text = str(h.get("local_ip", "NA") or "NA")
                    h_row[3].text = str(h.get("status", "NA") or "NA")
            else:
                doc.add_paragraph("No hosts found.")
                
            static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "reports")
            os.makedirs(static_dir, exist_ok=True)
            filename = f"hosts_and_groups_{int(time.time())}.docx"
            filepath = os.path.join(static_dir, filename)
            doc.save(filepath)
            
            download_links.append(f"**Generated Report:** [Download Hosts and Host Groups (DOCX)](/static/reports/{filename})")
            
        except Exception as e:
            logger.error(f"Failed to generate DOCX: {e}")
            prompt_context += f"\n\n[Error generating DOCX: {e}]\n"
            
    if generate_csv and fetch_host_groups:
        try:
            yield {
                "event": "agent_state",
                "data": {
                    "agent": "Policy Analyst",
                    "message": "Generating CSV report for Hosts and Host Groups..."
                }
            }
            
            static_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "static", "reports")
            os.makedirs(static_dir, exist_ok=True)
            filename_csv = f"hosts_and_groups_{int(time.time())}.csv"
            filepath_csv = os.path.join(static_dir, filename_csv)
            
            with open(filepath_csv, mode='w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                
                # Host Groups
                writer.writerow(["--- HOST GROUPS ---"])
                writer.writerow(["Group Name", "Group ID", "Hosts Applied", "Group Type", "Group Description", "Created By", "Created Date", "Last Modified By", "Last Modified", "Assignment Rule"])
                if host_groups:
                    for g in host_groups:
                        writer.writerow([
                            str(g.get("name", "NA") or "NA"),
                            str(g.get("id", "NA") or "NA"),
                            str(g.get("hosts_applied", "NA")),
                            str(g.get("group_type", "NA") or "NA"),
                            str(g.get("description", "NA") or "NA"),
                            str(g.get("created_by", "NA") or "NA"),
                            str(g.get("created_timestamp", "NA") or "NA"),
                            str(g.get("modified_by", "NA") or "NA"),
                            str(g.get("modified_timestamp", "NA") or "NA"),
                            str(g.get("assignment_rule", "NA") or "NA")
                        ])
                else:
                    writer.writerow(["No host groups found."])
                    
                writer.writerow([])
                writer.writerow([])
                
                # Hosts
                writer.writerow(["--- MANAGED HOSTS ---"])
                writer.writerow(["Hostname", "OS Version", "Local IP", "Status"])
                if hosts:
                    for h in (hosts or []):
                        writer.writerow([
                            str(h.get("hostname", "NA") or "NA"),
                            str(h.get("os_version", "NA") or "NA"),
                            str(h.get("local_ip", "NA") or "NA"),
                            str(h.get("status", "NA") or "NA")
                        ])
                else:
                    writer.writerow(["No hosts found."])
            
            download_links.append(f"**Generated CSV:** [Download Hosts and Host Groups (CSV)](/static/reports/{filename_csv})")
            
        except Exception as e:
            logger.error(f"Failed to generate CSV: {e}")
            prompt_context += f"\n\n[Error generating CSV: {e}]\n"
            
    # Call LLM stream
    try:
        sys_note = ""
        if download_links:
            sys_note = "\n\n[SYSTEM NOTE: The requested export file (DOCX/CSV) has ALREADY been generated successfully by the backend and the download link will be appended to your response automatically. DO NOT apologize or state that you cannot generate files. Simply summarize the host groups data provided below.]\n"
            
        user_prompt = f"User Query: {user_query}{sys_note}\n\nConfigurations context from Falcon API:\n{prompt_context}"
        token_stream = call_llm_stream(SYSTEM_PROMPT, user_prompt, chat_history)
        for token in token_stream:
            yield {
                "event": "text_chunk",
                "data": {"text": token}
            }
    except Exception as e:
        logger.error(f"Error in Policy Analyst LLM stream: {e}")
        yield {
            "event": "text_chunk",
            "data": {"text": f"\n[Error: Failed generating policy analysis. Detail: {str(e)}]\n"}
        }
        
    if download_links:
        links_text = "\n\n" + "\n\n".join(download_links) + "\n"
        yield {
            "event": "text_chunk",
            "data": {"text": links_text}
        }
        
    yield {
        "event": "complete",
        "data": {
            "agent": "Policy Analyst",
            "status": "success"
        }
    }
