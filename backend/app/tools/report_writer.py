"""
SOC Report Writer Module
=========================
Generates professional SOC reports in Word (.docx), Excel (.xlsx),
and PDF formats from real CrowdStrike Falcon API data.

Supports report periods: daily, weekly, monthly, quarterly, annual.
"""
import os
import logging
from datetime import datetime, timezone
from typing import Dict, Any, List, Optional
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

logger = logging.getLogger("report_writer")

PERIOD_LABELS = {
    "daily": "Daily",
    "weekly": "Weekly",
    "monthly": "Monthly",
    "quarterly": "Quarterly",
    "annual": "Annual",
}

SEVERITY_COLORS = {
    "CRITICAL": "C00000",
    "HIGH": "FF0000",
    "MEDIUM": "FFA500",
    "LOW": "FFFF00",
    "INFO": "00B0F0",
}


def generate_soc_report(
    period: str,
    data: Dict[str, Any],
    output_dir: str,
) -> Dict[str, str]:
    """
    Entry point for SOC report generation.
    Creates Word, Excel, and PDF reports for the specified period.

    Args:
        period: Report period — "daily", "weekly", "monthly", "quarterly", "annual"
        data: Assembled report data dict from the report_generator agent
        output_dir: Directory path to save generated report files

    Returns:
        Dict mapping format to file path: {"docx": "...", "xlsx": "...", "pdf": "..."}
    """
    os.makedirs(output_dir, exist_ok=True)
    period_label = PERIOD_LABELS.get(period.lower(), period.title())
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    base_name = f"soc_{period.lower()}_report_{timestamp}"

    paths = {
        "docx": os.path.join(output_dir, f"{base_name}.docx"),
        "xlsx": os.path.join(output_dir, f"{base_name}.xlsx"),
        "pdf": os.path.join(output_dir, f"{base_name}.pdf"),
        "html": os.path.join(output_dir, f"{base_name}.html"),
    }

    data["period_label"] = period_label

    try:
        generate_word_report(data, paths["docx"])
        logger.info(f"Word report generated: {paths['docx']}")
    except Exception as e:
        logger.error(f"Word report generation failed: {e}")
        paths["docx"] = None

    try:
        generate_excel_report(data, paths["xlsx"])
        logger.info(f"Excel report generated: {paths['xlsx']}")
    except Exception as e:
        logger.error(f"Excel report generation failed: {e}")
        paths["xlsx"] = None

    try:
        generate_pdf_report(data, paths["pdf"])
        logger.info(f"PDF report generated: {paths['pdf']}")
    except Exception as e:
        logger.error(f"PDF/HTML report generation failed: {e}")
        paths["pdf"] = paths.get("html")

    return {k: v for k, v in paths.items() if v}


def generate_word_report(data: Dict[str, Any], file_path: str):
    """
    Generates a professional Word (.docx) SOC report.
    """
    doc = Document()
    period_label = data.get("period_label", "Security")
    title = f"Falcon AI Copilot — {period_label} SOC Report"

    # Title
    title_para = doc.add_heading(title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata block
    meta_table = doc.add_table(rows=4, cols=2)
    meta_table.style = "Table Grid"
    meta_cells = [
        ("Report Period", data.get("period_label", "")),
        ("Generated At", data.get("timestamp", "")),
        ("Scope", data.get("scope", "Enterprise Endpoints")),
        ("Classification", "CONFIDENTIAL — Internal Use Only"),
    ]
    for i, (label, value) in enumerate(meta_cells):
        meta_table.rows[i].cells[0].text = label
        meta_table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    summary = data.get("executive_summary", "")
    if summary:
        doc.add_paragraph(summary)

    # === DETECTIONS SECTION ===
    doc.add_heading("Detection Summary", level=1)
    det_metrics = data.get("detections", {})
    total_det = det_metrics.get("total", "N/A")
    by_sev = det_metrics.get("by_severity", {})

    p = doc.add_paragraph()
    p.add_run(f"Total Detections: ").bold = True
    p.add_run(str(total_det))

    if by_sev:
        sev_table = doc.add_table(rows=1, cols=2)
        sev_table.style = "Table Grid"
        sev_table.rows[0].cells[0].text = "Severity"
        sev_table.rows[0].cells[1].text = "Count"
        for sev, count in by_sev.items():
            row = sev_table.add_row()
            row.cells[0].text = sev
            row.cells[1].text = str(count)

    doc.add_paragraph()

    # Detection list
    detections = data.get("detection_list", [])
    if detections:
        doc.add_heading("Detection Details", level=2)
        det_table = doc.add_table(rows=1, cols=5)
        det_table.style = "Table Grid"
        headers = ["Detection ID", "Host", "Severity", "Technique", "Created"]
        for i, h in enumerate(headers):
            det_table.rows[0].cells[i].text = h
        for det in detections[:50]:  # Limit to 50 for document size
            row = det_table.add_row()
            row.cells[0].text = str(det.get("detection_id", det.get("id", "")))
            row.cells[1].text = str(det.get("device", {}).get("hostname", det.get("host", "")))
            row.cells[2].text = str(det.get("max_severity_displayname", det.get("severity", "")))
            behaviors = det.get("behaviors", [{}])
            row.cells[3].text = str(behaviors[0].get("technique", "") if behaviors else "")
            row.cells[4].text = str(det.get("created_timestamp", ""))[:19]

    doc.add_paragraph()

    # === INCIDENTS SECTION ===
    doc.add_heading("Incident Summary", level=1)
    inc_metrics = data.get("incidents", {})
    total_inc = inc_metrics.get("total", "N/A")
    p = doc.add_paragraph()
    p.add_run("Total Incidents: ").bold = True
    p.add_run(str(total_inc))

    incidents = data.get("incident_list", [])
    if incidents:
        inc_table = doc.add_table(rows=1, cols=4)
        inc_table.style = "Table Grid"
        for i, h in enumerate(["Incident ID", "Status", "Score", "Created"]):
            inc_table.rows[0].cells[i].text = h
        for inc in incidents[:20]:
            row = inc_table.add_row()
            row.cells[0].text = str(inc.get("incident_id", inc.get("id", "")))
            row.cells[1].text = str(inc.get("status", ""))
            row.cells[2].text = str(inc.get("fine_score", ""))
            row.cells[3].text = str(inc.get("created_timestamp", ""))[:19]

    doc.add_paragraph()

    # === SENSOR COVERAGE SECTION ===
    doc.add_heading("Sensor Coverage", level=1)
    coverage = data.get("sensor_coverage", {})
    if coverage:
        cov_table = doc.add_table(rows=1, cols=2)
        cov_table.style = "Table Grid"
        cov_table.rows[0].cells[0].text = "Metric"
        cov_table.rows[0].cells[1].text = "Value"
        cov_items = [
            ("Total Managed Hosts", coverage.get("total_managed_hosts", "N/A")),
            ("Online (Last 24h)", coverage.get("online_last_24h", "N/A")),
            ("Offline (7+ days)", coverage.get("offline_7d", "N/A")),
            ("Reduced Functionality Mode", coverage.get("reduced_functionality_mode", "N/A")),
            ("Coverage %", f"{coverage.get('coverage_percentage', 'N/A')}%"),
        ]
        for label, value in cov_items:
            row = cov_table.add_row()
            row.cells[0].text = str(label)
            row.cells[1].text = str(value)

    doc.add_paragraph()

    # === VULNERABILITY SECTION ===
    doc.add_heading("Vulnerability Status (Falcon Spotlight)", level=1)
    vulns = data.get("vulnerabilities", {})
    if vulns:
        p = doc.add_paragraph()
        p.add_run("Total Open Vulnerabilities: ").bold = True
        p.add_run(str(vulns.get("total_open", "N/A")))
        p = doc.add_paragraph()
        p.add_run("CISA KEV Count: ").bold = True
        p.add_run(str(vulns.get("kev_count", "N/A")))

        sev_data = vulns.get("by_severity", {})
        if sev_data:
            vuln_table = doc.add_table(rows=1, cols=2)
            vuln_table.style = "Table Grid"
            vuln_table.rows[0].cells[0].text = "Severity"
            vuln_table.rows[0].cells[1].text = "Count"
            for sev, count in sev_data.items():
                row = vuln_table.add_row()
                row.cells[0].text = sev
                row.cells[1].text = str(count)

    doc.add_paragraph()

    # === KPI SECTION ===
    doc.add_heading("Key Performance Indicators", level=1)
    mttd = data.get("mttd", {})
    mttr = data.get("mttr", {})

    if mttd.get("mttd_human"):
        p = doc.add_paragraph()
        p.add_run("Mean Time to Detect (MTTD): ").bold = True
        p.add_run(str(mttd.get("mttd_human")))

    if mttr.get("mttr_human"):
        p = doc.add_paragraph()
        p.add_run("Mean Time to Respond (MTTR): ").bold = True
        p.add_run(str(mttr.get("mttr_human")))

    # Top Threats
    top_threats = data.get("top_threats", {})
    if top_threats.get("top_techniques"):
        doc.add_heading("Top MITRE Techniques Detected", level=2)
        for item in top_threats["top_techniques"][:10]:
            doc.add_paragraph(
                f"• {item['name']} — {item['count']} detections",
                style="List Bullet"
            )

    doc.add_paragraph()

    # Footer
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(
        "This report was generated automatically by Falcon AI Copilot operating in "
        "strict READ-ONLY mode. No CrowdStrike resources were modified. "
        "All data is sourced directly from the CrowdStrike Falcon platform APIs. "
        "CONFIDENTIAL — FOR INTERNAL USE ONLY."
    )

    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    doc.save(file_path)
    logger.info(f"Word report saved: {file_path}")


def generate_excel_report(data: Dict[str, Any], file_path: str):
    """
    Generates a professional Excel (.xlsx) SOC report with multiple worksheets.
    """
    wb = openpyxl.Workbook()
    period_label = data.get("period_label", "Security")
    now_str = data.get("timestamp", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))

    # ── Summary Sheet ────────────────────────────────────────────────────────
    ws_summary = wb.active
    ws_summary.title = "Executive Summary"
    _write_excel_header(ws_summary, f"Falcon AI Copilot — {period_label} SOC Report", now_str)

    row = 4
    ws_summary.cell(row=row, column=1, value="Metric").font = Font(bold=True)
    ws_summary.cell(row=row, column=2, value="Value").font = Font(bold=True)
    row += 1

    summary_rows = [
        ("Total Detections", data.get("detections", {}).get("total", "N/A")),
        ("Total Incidents", data.get("incidents", {}).get("total", "N/A")),
        ("Total Managed Hosts", data.get("sensor_coverage", {}).get("total_managed_hosts", "N/A")),
        ("Online Hosts (24h)", data.get("sensor_coverage", {}).get("online_last_24h", "N/A")),
        ("Offline Hosts (7d)", data.get("sensor_coverage", {}).get("offline_7d", "N/A")),
        ("Coverage %", f"{data.get('sensor_coverage', {}).get('coverage_percentage', 'N/A')}%"),
        ("Open Vulnerabilities", data.get("vulnerabilities", {}).get("total_open", "N/A")),
        ("CISA KEV Count", data.get("vulnerabilities", {}).get("kev_count", "N/A")),
        ("MTTD", data.get("mttd", {}).get("mttd_human", "N/A")),
        ("MTTR", data.get("mttr", {}).get("mttr_human", "N/A")),
    ]

    for label, value in summary_rows:
        ws_summary.cell(row=row, column=1, value=label)
        ws_summary.cell(row=row, column=2, value=str(value))
        row += 1

    ws_summary.column_dimensions["A"].width = 35
    ws_summary.column_dimensions["B"].width = 25

    # ── Detections Sheet ─────────────────────────────────────────────────────
    ws_det = wb.create_sheet("Detections")
    _write_excel_header(ws_det, f"Detections — {period_label}", now_str)

    det_headers = ["Detection ID", "Host", "Severity", "Tactic", "Technique", "Created", "Status"]
    for col, header in enumerate(det_headers, 1):
        cell = ws_det.cell(row=4, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F3864")

    for row_num, det in enumerate(data.get("detection_list", []), 5):
        behaviors = det.get("behaviors", [{}])
        b = behaviors[0] if behaviors else {}
        sev = det.get("max_severity_displayname", det.get("severity", ""))
        cells = [
            det.get("detection_id", det.get("id", "")),
            det.get("device", {}).get("hostname", det.get("host", "")),
            sev,
            b.get("tactic", ""),
            b.get("technique", ""),
            str(det.get("created_timestamp", ""))[:19],
            det.get("status", "new"),
        ]
        for col, val in enumerate(cells, 1):
            c = ws_det.cell(row=row_num, column=col, value=str(val))
            if col == 3 and sev in SEVERITY_COLORS:
                c.fill = PatternFill("solid", fgColor=SEVERITY_COLORS[sev])

    for col in range(1, len(det_headers) + 1):
        ws_det.column_dimensions[get_column_letter(col)].width = 22

    # ── Incidents Sheet ───────────────────────────────────────────────────────
    ws_inc = wb.create_sheet("Incidents")
    _write_excel_header(ws_inc, f"Incidents — {period_label}", now_str)

    inc_headers = ["Incident ID", "Status", "Fine Score", "Hosts", "Created", "Tactics"]
    for col, header in enumerate(inc_headers, 1):
        cell = ws_inc.cell(row=4, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F3864")

    for row_num, inc in enumerate(data.get("incident_list", []), 5):
        hosts_str = ", ".join([
            h.get("hostname", h) if isinstance(h, dict) else str(h)
            for h in inc.get("hosts", [])[:3]
        ])
        tactics_str = ", ".join(inc.get("tactics", []))
        cells = [
            inc.get("incident_id", inc.get("id", "")),
            inc.get("status", ""),
            inc.get("fine_score", ""),
            hosts_str,
            str(inc.get("created_timestamp", ""))[:19],
            tactics_str[:100],
        ]
        for col, val in enumerate(cells, 1):
            ws_inc.cell(row=row_num, column=col, value=str(val))

    for col in range(1, len(inc_headers) + 1):
        ws_inc.column_dimensions[get_column_letter(col)].width = 22

    # ── Vulnerabilities Sheet ─────────────────────────────────────────────────
    ws_vuln = wb.create_sheet("Vulnerabilities")
    _write_excel_header(ws_vuln, f"Vulnerabilities — {period_label}", now_str)

    vuln_headers = ["CVE ID", "Severity", "CVSS", "Host", "Status", "Is KEV"]
    for col, header in enumerate(vuln_headers, 1):
        cell = ws_vuln.cell(row=4, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F3864")

    for row_num, vuln in enumerate(data.get("vulnerability_list", []), 5):
        cve = vuln.get("cve", {})
        cells = [
            cve.get("id", ""),
            cve.get("severity", ""),
            cve.get("base_score", ""),
            vuln.get("host_info", {}).get("hostname", ""),
            vuln.get("status", ""),
            str(cve.get("is_kev", False)),
        ]
        sev = cve.get("severity", "")
        for col, val in enumerate(cells, 1):
            c = ws_vuln.cell(row=row_num, column=col, value=str(val))
            if col == 2 and sev in SEVERITY_COLORS:
                c.fill = PatternFill("solid", fgColor=SEVERITY_COLORS[sev])

    for col in range(1, len(vuln_headers) + 1):
        ws_vuln.column_dimensions[get_column_letter(col)].width = 20

    # ── Sensor Coverage Sheet ─────────────────────────────────────────────────
    ws_sensor = wb.create_sheet("Sensor Coverage")
    _write_excel_header(ws_sensor, f"Sensor Coverage — {period_label}", now_str)

    coverage = data.get("sensor_coverage", {})
    cov_rows = [
        ("Total Managed Hosts", coverage.get("total_managed_hosts", "N/A")),
        ("Online (Last 24h)", coverage.get("online_last_24h", "N/A")),
        ("Offline (7+ days)", coverage.get("offline_7d", "N/A")),
        ("Reduced Functionality Mode", coverage.get("reduced_functionality_mode", "N/A")),
        ("Coverage Percentage", f"{coverage.get('coverage_percentage', 'N/A')}%"),
    ]
    for row_num, (label, value) in enumerate(cov_rows, 4):
        ws_sensor.cell(row=row_num, column=1, value=label).font = Font(bold=True)
        ws_sensor.cell(row=row_num, column=2, value=str(value))

    ws_sensor.column_dimensions["A"].width = 35
    ws_sensor.column_dimensions["B"].width = 20

    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    wb.save(file_path)
    logger.info(f"Excel report saved: {file_path}")


def _write_excel_header(ws, title: str, timestamp: str):
    """Writes a standard header to an Excel worksheet."""
    ws.merge_cells("A1:G1")
    title_cell = ws["A1"]
    title_cell.value = title
    title_cell.font = Font(size=14, bold=True, color="FFFFFF")
    title_cell.fill = PatternFill("solid", fgColor="1F3864")
    title_cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    ws.merge_cells("A2:G2")
    ts_cell = ws["A2"]
    ts_cell.value = f"Generated: {timestamp} | Falcon AI Copilot | READ-ONLY"
    ts_cell.font = Font(size=10, italic=True, color="555555")
    ts_cell.alignment = Alignment(horizontal="center")


def generate_pdf_report(data: Dict[str, Any], file_path: str):
    """
    Generates a PDF report via WeasyPrint.
    Falls back to saving an HTML file if WeasyPrint is unavailable.
    """
    period_label = data.get("period_label", "Security")
    timestamp = data.get("timestamp", "")
    scope = data.get("scope", "Enterprise Endpoints")

    # Build severity table rows
    sev_rows = ""
    for sev, count in data.get("detections", {}).get("by_severity", {}).items():
        color = {
            "CRITICAL": "#c00000", "HIGH": "#ff4444",
            "MEDIUM": "#ff8c00", "LOW": "#ffd700"
        }.get(sev, "#666")
        sev_rows += f'<tr><td style="color:{color};font-weight:bold">{sev}</td><td>{count}</td></tr>'

    # Build detection table rows
    det_rows = ""
    for det in data.get("detection_list", [])[:30]:
        behaviors = det.get("behaviors", [{}])
        b = behaviors[0] if behaviors else {}
        sev = det.get("max_severity_displayname", det.get("severity", ""))
        color = {
            "CRITICAL": "#c00000", "HIGH": "#ff4444",
            "MEDIUM": "#ff8c00", "LOW": "#ffd700"
        }.get(sev, "#666")
        det_rows += f"""<tr>
            <td>{det.get("detection_id", det.get("id", ""))}</td>
            <td>{det.get("device", {}).get("hostname", det.get("host", ""))}</td>
            <td style="color:{color};font-weight:bold">{sev}</td>
            <td>{b.get("technique", "")}</td>
            <td>{str(det.get("created_timestamp", ""))[:19]}</td>
        </tr>"""

    # Sensor coverage data
    coverage = data.get("sensor_coverage", {})
    vuln = data.get("vulnerabilities", {})
    mttd = data.get("mttd", {}).get("mttd_human", "N/A")
    mttr = data.get("mttr", {}).get("mttr_human", "N/A")

    html_content = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<title>Falcon AI Copilot — {period_label} SOC Report</title>
<style>
  * {{ margin: 0; padding: 0; box-sizing: border-box; }}
  body {{ font-family: 'Segoe UI', Arial, sans-serif; color: #1a1a2e; background: #fff; padding: 40px; }}
  .report-header {{ background: linear-gradient(135deg, #1a1a2e, #16213e); color: white; padding: 30px; border-radius: 8px; margin-bottom: 30px; }}
  .report-header h1 {{ font-size: 24px; margin-bottom: 5px; }}
  .report-header p {{ font-size: 12px; opacity: 0.8; }}
  .kpi-grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 15px; margin-bottom: 30px; }}
  .kpi-card {{ background: #f8f9fa; border: 1px solid #dee2e6; border-radius: 8px; padding: 15px; text-align: center; }}
  .kpi-card .label {{ font-size: 11px; color: #6c757d; text-transform: uppercase; letter-spacing: 0.5px; }}
  .kpi-card .value {{ font-size: 28px; font-weight: bold; color: #1a1a2e; margin-top: 5px; }}
  .section {{ margin-bottom: 25px; }}
  .section h2 {{ font-size: 16px; color: #1a1a2e; border-bottom: 2px solid #e53e3e; padding-bottom: 5px; margin-bottom: 15px; }}
  table {{ width: 100%; border-collapse: collapse; font-size: 12px; }}
  th {{ background: #1a1a2e; color: white; padding: 8px 10px; text-align: left; }}
  td {{ padding: 7px 10px; border-bottom: 1px solid #dee2e6; }}
  tr:nth-child(even) {{ background: #f8f9fa; }}
  .badge-crit {{ background: #c00000; color: white; padding: 2px 8px; border-radius: 4px; font-size: 10px; font-weight: bold; }}
  .badge-high {{ background: #ff4444; color: white; padding: 2px 8px; border-radius: 4px; font-size: 10px; font-weight: bold; }}
  .badge-med {{ background: #ff8c00; color: white; padding: 2px 8px; border-radius: 4px; font-size: 10px; font-weight: bold; }}
  .badge-low {{ background: #ffd700; color: #333; padding: 2px 8px; border-radius: 4px; font-size: 10px; font-weight: bold; }}
  .footer {{ margin-top: 40px; font-size: 10px; color: #6c757d; border-top: 1px solid #dee2e6; padding-top: 15px; text-align: center; }}
</style>
</head>
<body>

<div class="report-header">
  <h1>🦅 Falcon AI Copilot — {period_label} SOC Report</h1>
  <p>Generated: {timestamp} | Scope: {scope} | READ-ONLY Mode</p>
</div>

<div class="kpi-grid">
  <div class="kpi-card">
    <div class="label">Total Detections</div>
    <div class="value">{data.get("detections", {}).get("total", "—")}</div>
  </div>
  <div class="kpi-card">
    <div class="label">Total Incidents</div>
    <div class="value">{data.get("incidents", {}).get("total", "—")}</div>
  </div>
  <div class="kpi-card">
    <div class="label">Sensor Coverage</div>
    <div class="value">{coverage.get("coverage_percentage", "—")}%</div>
  </div>
  <div class="kpi-card">
    <div class="label">Open Vulns</div>
    <div class="value">{vuln.get("total_open", "—")}</div>
  </div>
  <div class="kpi-card">
    <div class="label">MTTD</div>
    <div class="value" style="font-size:16px">{mttd}</div>
  </div>
  <div class="kpi-card">
    <div class="label">MTTR</div>
    <div class="value" style="font-size:16px">{mttr}</div>
  </div>
  <div class="kpi-card">
    <div class="label">KEV Count</div>
    <div class="value">{vuln.get("kev_count", "—")}</div>
  </div>
  <div class="kpi-card">
    <div class="label">Offline Sensors</div>
    <div class="value">{coverage.get("offline_7d", "—")}</div>
  </div>
</div>

<div class="section">
  <h2>Detection Severity Breakdown</h2>
  <table>
    <tr><th>Severity</th><th>Detection Count</th></tr>
    {sev_rows}
  </table>
</div>

<div class="section">
  <h2>Recent Detections</h2>
  <table>
    <tr><th>Detection ID</th><th>Host</th><th>Severity</th><th>Technique</th><th>Created</th></tr>
    {det_rows}
  </table>
</div>

<div class="section">
  <h2>Sensor Coverage</h2>
  <table>
    <tr><th>Metric</th><th>Value</th></tr>
    <tr><td>Total Managed Hosts</td><td>{coverage.get("total_managed_hosts", "N/A")}</td></tr>
    <tr><td>Online (Last 24h)</td><td>{coverage.get("online_last_24h", "N/A")}</td></tr>
    <tr><td>Offline (7+ days)</td><td>{coverage.get("offline_7d", "N/A")}</td></tr>
    <tr><td>Reduced Functionality Mode</td><td>{coverage.get("reduced_functionality_mode", "N/A")}</td></tr>
  </table>
</div>

<div class="footer">
  CONFIDENTIAL — Internal Use Only · Falcon AI Copilot · Strictly READ-ONLY
</div>
</body>
</html>"""

    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    try:
        from weasyprint import HTML
        HTML(string=html_content).write_pdf(file_path)
        logger.info(f"PDF report saved via WeasyPrint: {file_path}")
    except Exception as e:
        logger.warning(f"WeasyPrint unavailable: {e}. Saving as HTML.")
        html_path = file_path.replace(".pdf", ".html")
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        logger.info(f"HTML fallback report saved: {html_path}")
